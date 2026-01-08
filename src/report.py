from datetime import datetime

from docx import Document


def generate_report(duplicates, output_path):
    print("Gerando relatório em...")
    doc = Document()
    doc.add_heading("Relatório de Cartas Duplicadas", level=1)
    doc.add_paragraph(f"Data da Análise: {datetime.now().strftime('%d/%m/Y %H:%M')}")

    if not duplicates:
        doc.add_paragraph("Nenhum duplicata encontrada.")
    else:
        table = doc.add_table(rows=1, cols=4)
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = "Setor"
        hdr_cells[1].text = "Código"
        hdr_cells[2].text = "Quantidade"
        hdr_cells[3].text = "Arquivos"

        for key, files in duplicates.items():
            setor, codigo = key.split("_")
            row_cells = table.add_row().cells
            row_cells[0].text = setor
            row_cells[1].text = codigo
            row_cells[2].text = str(len(files))
            row_cells[3].text = ", ".join(files)

    doc.save(output_path)
