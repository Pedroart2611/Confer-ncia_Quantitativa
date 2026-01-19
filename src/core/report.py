from datetime import datetime
from typing import Dict, List
from docx import Document

def generate_report(duplicates: Dict[str, List[str]], output_path: str) -> None:
    '''
    Gera um relatório em um documento Word (.docx) com os arquivos escaneados e suas duplicatas

    parametros:
        duplicates: Dicionário onde a chave é 'setor_codigo' e o valor é a lista de arquivos
        output_path:  Caminho onde o relatório será salvo
    '''

    print("Gerando relatório em... ")
    document = Document()
    document.add_heading("Relatório das cartas duplicadas", level=1)
    document.add_paragraph(f"Data da conferência: {datetime.now().strftime('%d/%m/%Y %H:%M:%S')}")

    if not duplicates:
        document.add_paragraph("Nenhuma duplicata encontrada.")
    else:
        table = document.add_table(rows=1, cols=4)
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = "Setor"
        hdr_cells[1].text = "Código"
        hdr_cells[2].text = "Quantidade"
        hdr_cells[3].text = "Arquivos"

        for key, files in duplicates.items():
            setor, codigo = key.split("_", 1)
            row_cells = table.add_row().cells
            row_cells[0].text = setor
            row_cells[1].text = codigo
            row_cells[2].text = str(len(files))
            row_cells[3].text = ", ".join(files)

    document.save(output_path)